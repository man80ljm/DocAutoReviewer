from __future__ import annotations

from pathlib import Path
from typing import Iterable, Optional, Tuple, Union

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


REFLECTION_MARKER = "[[REFLECTION_CELL]]"
COMMENT_MARKER = "[[COMMENT_HERE]]"
REFLECTION_KEYWORDS = [
    "我的实验心得",
    "实验心得",
    "心得与反思",
    "Student's Reflection",
]


def extract_reflection_text(
    doc_path: Path,
    max_chars: int = 4000,
    start_marker: Optional[str] = None,
    end_marker: Optional[str] = None,
) -> Optional[str]:
    doc = Document(str(doc_path))
    marker_text = _find_reflection_by_range_markers(doc, start_marker, end_marker)
    if marker_text:
        return _clean_text(marker_text, max_chars)
    marker_text = _find_reflection_by_marker(doc)
    if marker_text:
        return _clean_text(marker_text, max_chars)
    fallback_text = _find_reflection_by_keywords(doc)
    if fallback_text:
        return _clean_text(fallback_text, max_chars)
    return None


def insert_comment_and_save(
    doc_path: Path,
    comment_text: str,
    output_dir: Path,
    comment_label: Optional[str] = None,
    overwrite_output: bool = True,
) -> Optional[Path]:
    doc = Document(str(doc_path))
    inserted = _insert_comment(doc, comment_text, comment_label)
    if not inserted:
        return None
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / doc_path.name if overwrite_output else _next_available_path(output_dir, doc_path.name)
    doc.save(str(output_path))
    return output_path


def _find_reflection_by_range_markers(
    doc: Document,
    start_marker: Optional[str],
    end_marker: Optional[str],
) -> Optional[str]:
    if not start_marker or not end_marker:
        return None
    collecting = False
    parts: list[str] = []
    for block in _iter_block_items(doc):
        block_text = _block_text(block).strip()
        if not block_text:
            continue
        if not collecting:
            if start_marker in block_text:
                collecting = True
                after = block_text.split(start_marker, 1)[1].strip()
                if end_marker in after:
                    before = after.split(end_marker, 1)[0].strip()
                    if before:
                        parts.append(before)
                    break
                if after:
                    parts.append(after)
        else:
            if end_marker in block_text:
                before = block_text.split(end_marker, 1)[0].strip()
                if before:
                    parts.append(before)
                break
            parts.append(block_text)
    if parts:
        return "\n".join(parts)
    return None


def _find_reflection_by_marker(doc: Document) -> Optional[str]:
    for cell in _iter_cells(doc):
        text = cell.text or ""
        if REFLECTION_MARKER in text:
            return text.replace(REFLECTION_MARKER, "")
    return None


def _find_reflection_by_keywords(doc: Document) -> Optional[str]:
    for table in doc.tables:
        hit = _find_keyword_cell(table)
        if hit:
            label_cell, row_idx, col_idx = hit
            # Try right cell first
            if col_idx + 1 < len(table.rows[row_idx].cells):
                right_cell = table.rows[row_idx].cells[col_idx + 1]
                right_text = right_cell.text.strip()
                if right_text:
                    return right_text
            # Try below cell
            if row_idx + 1 < len(table.rows):
                below_cell = table.rows[row_idx + 1].cells[col_idx]
                below_text = below_cell.text.strip()
                if below_text:
                    return below_text
            # Fallback to same cell content, removing keyword line
            cell_text = label_cell.text
            for kw in REFLECTION_KEYWORDS:
                cell_text = cell_text.replace(kw, "")
            if cell_text.strip():
                return cell_text
    return None


def _find_keyword_cell(table: Table) -> Optional[Tuple[_Cell, int, int]]:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text or ""
            for kw in REFLECTION_KEYWORDS:
                if kw in text:
                    return cell, r_idx, c_idx
    return None


def _insert_comment(doc: Document, comment_text: str, comment_label: Optional[str]) -> bool:
    if _replace_comment_marker(doc, comment_text):
        return True
    label = comment_label or "教师总评"
    return _insert_after_teacher_label(doc, comment_text, label)


def _replace_comment_marker(doc: Document, comment_text: str) -> bool:
    replaced = False
    for para in doc.paragraphs:
        if COMMENT_MARKER in para.text:
            para.text = para.text.replace(COMMENT_MARKER, comment_text)
            replaced = True
    for cell in _iter_cells(doc):
        for para in cell.paragraphs:
            if COMMENT_MARKER in para.text:
                para.text = para.text.replace(COMMENT_MARKER, comment_text)
                replaced = True
    return replaced


def _insert_after_teacher_label(doc: Document, comment_text: str, label: str) -> bool:
    # Paragraphs first
    for para in doc.paragraphs:
        if label in para.text:
            if para.text.strip() == label:
                para.text = f"{label}\n{comment_text}"
            else:
                para.text = para.text.replace(label, f"{label}\n{comment_text}", 1)
            return True
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for c_idx, cell in enumerate(row.cells):
                if label in (cell.text or ""):
                    # Prefer right cell for content
                    if c_idx + 1 < len(row.cells):
                        target = row.cells[c_idx + 1]
                        if target.paragraphs:
                            target.paragraphs[0].text = comment_text
                        else:
                            target.add_paragraph(comment_text)
                        return True
                    # Otherwise same cell
                    if cell.paragraphs:
                        cell.paragraphs[0].text = cell.paragraphs[0].text.replace(
                            label, f"{label}\n{comment_text}", 1
                        )
                    else:
                        cell.add_paragraph(comment_text)
                    return True
    return False


def _iter_cells(doc: Document) -> Iterable[_Cell]:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell


def _iter_block_items(parent: Union[DocxDocument, _Cell]) -> Iterable[Union[Paragraph, Table]]:
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _block_text(block: Union[Paragraph, Table]) -> str:
    if isinstance(block, Paragraph):
        return block.text or ""
    if isinstance(block, Table):
        parts: list[str] = []
        for row in block.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)
        return "\n".join(parts)
    return ""


def _clean_text(text: str, max_chars: int) -> str:
    normalized = " ".join(text.replace("\r", "\n").split())
    if len(normalized) > max_chars:
        return normalized[:max_chars]
    return normalized


def _next_available_path(output_dir: Path, filename: str) -> Path:
    base = Path(filename)
    candidate = output_dir / base.name
    if not candidate.exists():
        return candidate
    stem = base.stem
    suffix = base.suffix
    idx = 2
    while True:
        candidate = output_dir / f"{stem}_v{idx}{suffix}"
        if not candidate.exists():
            return candidate
        idx += 1
